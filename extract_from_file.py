'''
extract chaque ligne du faux fichier et générer un fichier pour chaque ligne
chaque fichier généré a en 5e ligne la ligné extraite du fichier
'''
from docx import Document
import streamlit as st
from io import StringIO, BytesIO
import time
from stqdm import stqdm
from zipfile import ZipFile


st.write("# Application de traitement de fichiers")
st.info("Chargez un fichier pour créer un fichier pour chaque ligne du fichier chargé. La ligne sera écrite en ligne 5 et sera centrée.")
fake_file = st.file_uploader("Chargez un fichier", type=["txt"])

def create_docs_in_memory(doc):
    
    in_memory = BytesIO()
    doc.save(in_memory)
    in_memory.seek(0)
    
    return in_memory
    
if fake_file:
    stringio = StringIO(fake_file.getvalue().decode("utf-8"))
    
    if st.button("Extraire"):
        lines = stringio.readlines()
        a = time.time()
        with ZipFile('fichiers.zip', 'w') as zipf:
            for i in stqdm(range(len(lines))):
                doc = Document()
                # put the 4 first empty lines
                for k in range(4):
                    doc.add_paragraph("")
                # put the extracted line on the 5th line of the document
                paragraph = doc.add_paragraph(lines[i])
                paragraph.alignment = 1
                doc_in_memory = create_docs_in_memory(doc)
                zipf.writestr(f"fichier_{i+1}.docx", doc_in_memory.getvalue())
        
        st.success("Fichiers crées et stocké dans un zip")
        st.balloons()
        b = time.time()
        delta = b - a
        st.write(f"Temps d'exécution : {round(delta, 2)} secondes pour extraire {i+1} lignes")
        
        with open('fichiers.zip', 'rb') as file:
            zip_content = file.read()
            st.download_button(
                label="Télécharger le fichier ZIP",
                data=zip_content,
                file_name='fichiers.zip',
                mime="application/zip"
                )
